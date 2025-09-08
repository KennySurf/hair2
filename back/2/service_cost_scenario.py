from services.gpt.gpt_client import send_to_gpt
from docx import Document
from db.db_funcs import update_services_id, update_state, get_user_messages, add_message
from services.gpt.gpt_client import send_to_gpt
from db.price_question_db_funcs import get_price_question_state, update_price_question_state, reset_price_question_state

def idle(user_id, user_messages,):
    print('idleee')

    prompt = """
    Правила (приоритет сверху вниз):
    
    1. Если имя клиента уже известно — выведи строго: знаю имя.
    
    2. Если пользователь задаёт вопрос про услуги (цены, процедуры, материалы, процесс и т.п.) —
       ответ должен содержать ровно две части:
       - "Сейчас всё расскажу"
       - если имя клиента ещё неизвестно — добавь только: "Как к вам обращаться?"
    
    3. Если вы ещё не здоровались за диалог — перед этими фразами добавь приветствие и представление в виде:
       "Добрый день! Меня зовут Мия, я менеджер по работе с клиентами студии EKSHE 🩷".
    
    Запрещено:
    — упоминать цены, консультацию, описание услуг;
    — задавать другие вопросы;
    — добавлять что-то помимо разрешённых фраз.
               """

    idle_reply = send_to_gpt(user_messages + [{'role': 'system', 'content': prompt}])

    if idle_reply == 'знаю имя':
        update_price_question_state(user_id, 'get_hook')
        print('knowww')
        return get_hook(user_id, user_messages)
    update_price_question_state(user_id, 'get_hook')
    print('not knoww')
    return idle_reply

def get_hook(user_id, user_messages):
    file = f'classifire_logic/question/files/общее положение.docx'

    doc = Document(file)
    full_text = "\n".join([para.text for para in doc.paragraphs])

    prompt = f"""
       Вежливо сообщи клиенту, что цену сказать можем только после консультации.
       Коротко одним предложением объясни от чего зависит цена - обьяснение можешь брать из доп информации.
       Если клиент спрашивает про наращивание - спроси его, определился ли он уже с методом наращивания - варианты наращивания возьми из доп информации.
       Если клиент спрашивает про окрашивание - спроси его, определился ли он уже с цветом.

       доп информация:
       {full_text}
       """

    update_price_question_state(user_id, 'get_visual')
    hook_reply = send_to_gpt(user_messages + [{'role': 'system', 'content': prompt}])
    return hook_reply

def get_visual(user_id, user_messages):
    prompt = f"""
    Оцени задал ли клиент какой то вопрос, или он обозначил свои желания
    
    Если он обозначил желания - определился с цветом/методом и т д - верни: определился
    Если вопрос - верни: вопрос
    Ничего лишнего не возвращай.
    """

    vis_class = send_to_gpt(user_messages + [{'role': 'system', 'content': prompt}])
    if vis_class == 'определился':
        service_prompt = f"""
        выдели услугу из запроса клиента.
        Наш пул:
        окрашивание
        наращивание

        Строго напиши только название нашей услуги, к которой подходит клиентский запрос.
        """

        service_name = send_to_gpt(
            user_messages[:-1] + [{'role': 'assistant', 'content': service_prompt}] + user_messages[-1:])

        if service_name == 'окрашивание':
            update_services_id(user_id, '17637553')
        else:
            update_services_id(user_id, '10928000')
        update_state(user_id, 'get_advice')

        agree_prompt = f"""Предложи визуализировать клиенту результат"""
        agree_reply = send_to_gpt(user_messages + [{'role': 'system', 'content': agree_prompt}])
        return agree_reply
    else:
        file = f'classifire_logic/question/files/общее положение.docx'

        doc = Document(file)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        other_prompt = f"""вежливо и кратко ответь на вопросы клиента пользуясь доп информацией
        доп информация: {full_text}
        """

        other_reply = send_to_gpt([{'role': 'system', 'content': other_prompt}])
        return other_reply
